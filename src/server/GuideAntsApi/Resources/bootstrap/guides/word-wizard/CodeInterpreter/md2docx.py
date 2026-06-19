# md2docx.py
"""
Library function to convert a Markdown (.md) file to a styled DOCX (.docx) file using a Word template.
This version fixes duplicate list item artifacts by only processing <li> tags for lists.

Usage Example:
from md2docx import md_to_docx
md_to_docx('input.md', 'template.docx', 'output.docx')

Arguments:
- md_path: Path to the Markdown file
- template_docx_path: Path to the DOCX template (should contain built-in Word styles)
- output_docx_path: Path for the output DOCX file
"""

def md_to_docx(md_path, template_docx_path, output_docx_path):
    import markdown2
    from bs4 import BeautifulSoup
    from docx import Document
    from docx.shared import Pt
    import os

    def has_style(doc, style_name):
        try:
            doc.styles[style_name]
            return True
        except KeyError:
            return False

    doc = Document(template_docx_path)
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()
    html_body = markdown2.markdown(md_text, extras=['fenced-code-blocks', 'tables'])
    soup = BeautifulSoup(html_body, 'html.parser')

    para_style = 'Normal' if has_style(doc, 'Normal') else None
    list_style = 'List Paragraph' if has_style(doc, 'List Paragraph') else None
    quote_style = 'Quote' if has_style(doc, 'Quote') else None
    strong_style = 'Strong' if has_style(doc, 'Strong') else None
    emph_style = 'Emphasis' if has_style(doc, 'Emphasis') else None

    # Only process top-level tags to avoid duplicate list item text
    for elem in soup.find_all(['h1', 'h2', 'h3', 'h4', 'p', 'pre', 'ul', 'ol', 'blockquote', 'table'], recursive=False):
        if elem.name == 'h1':
            doc.add_heading(elem.get_text(), level=1)
        elif elem.name == 'h2':
            doc.add_heading(elem.get_text(), level=2)
        elif elem.name == 'h3':
            doc.add_heading(elem.get_text(), level=3)
        elif elem.name == 'h4':
            doc.add_heading(elem.get_text(), level=4)
        elif elem.name == 'p':
            doc.add_paragraph(elem.get_text(), style=para_style)
        elif elem.name == 'pre':
            p = doc.add_paragraph()
            run = p.add_run(elem.get_text())
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif elem.name in ['ul', 'ol']:
            for li in elem.find_all('li', recursive=False):
                doc.add_paragraph(li.get_text(), style=list_style)
        elif elem.name == 'blockquote':
            doc.add_paragraph(elem.get_text(), style=quote_style)
        elif elem.name == 'table':
            rows = elem.find_all('tr')
            if rows:
                cols = rows[0].find_all(['td', 'th'])
                table = doc.add_table(rows=len(rows), cols=len(cols))
                for i, row in enumerate(rows):
                    cells = row.find_all(['td', 'th'])
                    for j, cell in enumerate(cells):
                        table.cell(i, j).text = cell.get_text()

    doc.save(output_docx_path)
    print(f"DOCX file created using template and styled: {output_docx_path}")
